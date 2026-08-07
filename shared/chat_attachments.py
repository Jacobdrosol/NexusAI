from __future__ import annotations

import base64
import io
import re
from typing import Any

CHAT_ATTACHMENT_MAX_FILES = 15
CHAT_ATTACHMENT_MAX_TOTAL_BYTES = 1024 * 1024 * 1024
CHAT_ATTACHMENT_MAX_TEXT_BYTES = 120_000
# Raw data URLs are persisted with the chat message. Keep this bounded so a single
# attachment cannot exhaust dashboard, API, or SQLite memory during a chat turn.
CHAT_ATTACHMENT_MAX_INLINE_BYTES = 25 * 1024 * 1024

_DATA_URL_RE = re.compile(r"^data:([^;,]+);base64,([A-Za-z0-9+/=]+)$", re.IGNORECASE)
_PDF_MIME_TYPES = {"application/pdf"}
_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def decode_attachment_data_url(value: Any, *, max_bytes: int = CHAT_ATTACHMENT_MAX_INLINE_BYTES) -> tuple[str, bytes]:
    """Decode a bounded base64 data URL without allowing executable URL schemes."""
    match = _DATA_URL_RE.match(str(value or "").strip())
    if not match:
        raise ValueError("attachment must provide a valid base64 data URL")
    try:
        raw = base64.b64decode(match.group(2), validate=True)
    except (ValueError, TypeError) as exc:
        raise ValueError("attachment data URL contains invalid base64") from exc
    if len(raw) > max_bytes:
        raise ValueError(f"attachment exceeds {max_bytes} byte inline limit")
    return match.group(1).lower(), raw


def extract_document_text(
    *,
    name: str,
    mime_type: str,
    raw: bytes,
    max_chars: int = CHAT_ATTACHMENT_MAX_TEXT_BYTES,
) -> tuple[str, str]:
    """Extract readable PDF or DOCX text for model context and attachment viewing."""
    normalized_mime = str(mime_type or "").strip().lower()
    lowered_name = str(name or "").strip().lower()
    try:
        if normalized_mime in _PDF_MIME_TYPES or lowered_name.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            if reader.is_encrypted:
                return "", "encrypted"
            chunks: list[str] = []
            remaining = max(1, int(max_chars))
            for page in reader.pages:
                text = str(page.extract_text() or "")
                if text:
                    chunks.append(text[:remaining])
                    remaining -= len(chunks[-1])
                if remaining <= 0:
                    break
            return "\n".join(chunks).strip(), "extracted"
        if normalized_mime in _DOCX_MIME_TYPES or lowered_name.endswith(".docx"):
            from docx import Document

            document = Document(io.BytesIO(raw))
            chunks = [str(paragraph.text or "").strip() for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    cells = [str(cell.text or "").strip() for cell in row.cells]
                    if any(cells):
                        chunks.append(" | ".join(cells))
            return "\n".join(item for item in chunks if item)[:max_chars], "extracted"
    except Exception:
        return "", "unavailable"
    return "", "unsupported"
