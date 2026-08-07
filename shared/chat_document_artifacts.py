"""Bounded DOCX artifact generation for explicitly enabled chat bots."""
from __future__ import annotations

import base64
import io
import json
import re
from copy import deepcopy
from typing import Any

from docx import Document

from shared.chat_attachments import CHAT_ATTACHMENT_MAX_INLINE_BYTES, CHAT_ATTACHMENT_MAX_TEXT_BYTES

_DOCX_REQUEST_RE = re.compile(r"(?:\.docx\b|\bword\s+document\b|\bmicrosoft\s+word\b)", re.IGNORECASE)
_DOCX_FILENAME_RE = re.compile(
    r"(?:\bnamed\b|\bcalled\b|\bas\b)\s+[\"']?([A-Za-z0-9][^\"'\n.?!]{0,95}\.docx)\b",
    re.IGNORECASE,
)
_INVALID_FILENAME_CHARS = re.compile(r"[^A-Za-z0-9._ -]+")
_DOCX_EDIT_REQUEST_RE = re.compile(
    r"\b(?:edit|tailor|revise|update|modify)\b.*\b(?:resume|document|docx|word)\b|"
    r"\b(?:preserve|keep|retain)\b.*\b(?:format|formatting|layout|styles?)\b",
    re.IGNORECASE,
)
_DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def document_generation_enabled(bot: Any) -> bool:
    """Return whether this bot is explicitly allowed to create chat DOCX artifacts."""
    routing_rules = getattr(bot, "routing_rules", None)
    if not isinstance(routing_rules, dict):
        return False
    profile = routing_rules.get("chat_profile")
    return bool(isinstance(profile, dict) and profile.get("document_generation") is True)


def document_editing_enabled(bot: Any) -> bool:
    """Return whether this bot may edit an attached DOCX while retaining its layout."""
    routing_rules = getattr(bot, "routing_rules", None)
    if not isinstance(routing_rules, dict):
        return False
    profile = routing_rules.get("chat_profile")
    return bool(isinstance(profile, dict) and profile.get("document_editing") is True)


def requested_docx_edit(
    prompt: str,
    attachments: list[dict[str, Any]],
    *,
    enabled: bool,
) -> tuple[dict[str, Any], str] | None:
    """Select one current-turn DOCX as the source for an explicit formatting-preserving edit."""
    if not enabled or not _DOCX_EDIT_REQUEST_RE.search(str(prompt or "")):
        return None
    source = next((item for item in attachments if _is_editable_docx_attachment(item)), None)
    if source is None:
        return None
    requested_name = _requested_filename(prompt)
    source_name = _safe_docx_filename(str(source.get("name") or "NexusAI_Document.docx"))
    if requested_name:
        output_name = requested_name
    else:
        stem = source_name[:-5]
        output_name = _safe_docx_filename(f"{stem}_edited.docx")
    return source, output_name


def requested_docx_artifact(prompt: str, *, enabled: bool) -> str | None:
    """Return the requested DOCX filename only for explicit, bot-authorized requests."""
    if not enabled or not _DOCX_REQUEST_RE.search(str(prompt or "")):
        return None
    return _requested_filename(prompt) or "NexusAI_Document.docx"


def document_generation_instruction(filename: str) -> str:
    return (
        "The user explicitly requested a downloadable DOCX named "
        f"'{filename}'. Provide the complete document body in Markdown now. "
        "Do not say that you cannot create files, do not include implementation notes, and do not wrap the "
        "document in a code fence. NexusAI will convert your final response into the DOCX attachment."
    )


def document_editing_instruction(*, filename: str, source_name: str) -> str:
    """Request a machine-applicable paragraph patch instead of a rebuilt document."""
    return (
        "The user asked NexusAI to edit the attached DOCX while preserving its formatting. "
        f"The source file is '{source_name}' and the returned copy will be named '{filename}'. "
        "Return ONLY valid JSON, with no Markdown fence or explanatory prose, in this exact shape: "
        '{"edits":[{"target":"exact existing paragraph text","replacement":"new paragraph text"}],'
        '"summary":"short description"}. '
        "Each target must exactly match one existing source paragraph supplied in the attachment context. "
        "Use replacement text for the entire paragraph, keep edits limited to the user's request, and do not invent "
        "experience, metrics, certifications, technologies, employers, or credentials. "
        "Do not edit headers, footers, contact details, section titles, or document layout unless the user explicitly asks."
    )


def build_docx_attachment(*, filename: str, content: str) -> dict[str, Any] | None:
    """Convert bounded Markdown-like assistant output into a downloadable DOCX attachment."""
    text = str(content or "").strip()
    if not text:
        return None
    document = Document()
    in_code_block = False
    for raw_line in text.splitlines():
        line = str(raw_line or "").rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if not stripped:
            document.add_paragraph("")
            continue
        if in_code_block:
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["No Spacing"]
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            document.add_heading(heading.group(2).strip(), level=len(heading.group(1)))
            continue
        if re.match(r"^[-*+]\s+", stripped):
            document.add_paragraph(re.sub(r"^[-*+]\s+", "", stripped), style="List Bullet")
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
            continue
        document.add_paragraph(stripped)

    buffer = io.BytesIO()
    document.save(buffer)
    raw = buffer.getvalue()
    if len(raw) > CHAT_ATTACHMENT_MAX_INLINE_BYTES:
        return None
    encoded = base64.b64encode(raw).decode("ascii")
    return {
        "name": _safe_docx_filename(filename),
        "mime_type": _DOCX_MIME_TYPE,
        "kind": "document",
        "data_url": f"data:{_DOCX_MIME_TYPE};base64,{encoded}",
        "text_content": text[:CHAT_ATTACHMENT_MAX_TEXT_BYTES],
        "size_bytes": len(raw),
        "truncated": len(text) > CHAT_ATTACHMENT_MAX_TEXT_BYTES,
        "extraction_status": "generated",
        "generated_by": "chat_document_artifact",
    }


def build_edited_docx_attachment(
    *,
    source_attachment: dict[str, Any],
    filename: str,
    model_output: str,
) -> tuple[dict[str, Any] | None, int, str]:
    """Apply exact paragraph replacements to a DOCX copy without rebuilding its structure.

    Paragraph-level properties, including Word list numbering and indentation, remain in
    place because only paragraph text is replaced. Complex inline run formatting inside
    an edited paragraph is intentionally not guaranteed; the untouched document retains
    its original OOXML unchanged.
    """
    plan = _parse_document_edit_plan(model_output)
    if plan is None:
        return None, 0, "The model did not return a valid document edit plan."
    data_url = str(source_attachment.get("data_url") or "").strip()
    try:
        from shared.chat_attachments import decode_attachment_data_url, extract_document_text

        mime_type, raw = decode_attachment_data_url(data_url)
        if not _is_docx_mime_or_name(mime_type, str(source_attachment.get("name") or "")):
            return None, 0, "The selected attachment is not a DOCX document."
        document = Document(io.BytesIO(raw))
    except Exception:
        return None, 0, "The attached DOCX could not be opened for editing."

    paragraphs = list(_iter_document_paragraphs(document))
    indexed: dict[str, list[Any]] = {}
    for paragraph in paragraphs:
        key = str(paragraph.text or "").strip()
        if key:
            indexed.setdefault(key, []).append(paragraph)

    changed = 0
    used_targets: set[tuple[str, int]] = set()
    for edit in plan["edits"]:
        target = str(edit["target"]).strip()
        replacement = str(edit["replacement"]).strip()
        candidates = indexed.get(target) or []
        paragraph = next(
            (candidate for candidate in candidates if (target, id(candidate)) not in used_targets),
            None,
        )
        if paragraph is None or not replacement:
            continue
        _replace_paragraph_text(paragraph, replacement)
        used_targets.add((target, id(paragraph)))
        changed += 1

    if changed == 0:
        return None, 0, "The edit plan did not match any source paragraphs exactly."

    buffer = io.BytesIO()
    document.save(buffer)
    edited_raw = buffer.getvalue()
    if len(edited_raw) > CHAT_ATTACHMENT_MAX_INLINE_BYTES:
        return None, 0, "The edited DOCX exceeds the chat attachment size limit."
    text_content, extraction_status = extract_document_text(
        name=filename,
        mime_type=_DOCX_MIME_TYPE,
        raw=edited_raw,
    )
    encoded = base64.b64encode(edited_raw).decode("ascii")
    return (
        {
            "name": _safe_docx_filename(filename),
            "mime_type": _DOCX_MIME_TYPE,
            "kind": "document",
            "data_url": f"data:{_DOCX_MIME_TYPE};base64,{encoded}",
            "text_content": text_content[:CHAT_ATTACHMENT_MAX_TEXT_BYTES],
            "size_bytes": len(edited_raw),
            "truncated": len(text_content) > CHAT_ATTACHMENT_MAX_TEXT_BYTES,
            "extraction_status": extraction_status,
            "generated_by": "chat_document_editor",
            "source_attachment_name": str(source_attachment.get("name") or "document.docx"),
            "preserved_source_formatting": True,
            "applied_edit_count": changed,
        },
        changed,
        str(plan.get("summary") or "").strip(),
    )


def _safe_docx_filename(value: str) -> str:
    name = _INVALID_FILENAME_CHARS.sub("", str(value or "").strip()).strip(". ")
    if not name:
        name = "NexusAI_Document.docx"
    if not name.lower().endswith(".docx"):
        name = f"{name}.docx"
    stem = name[:-5].strip()[:100].rstrip(". ") or "NexusAI_Document"
    return f"{stem}.docx"


def _requested_filename(prompt: str) -> str | None:
    match = _DOCX_FILENAME_RE.search(str(prompt or ""))
    return _safe_docx_filename(match.group(1)) if match else None


def _is_docx_mime_or_name(mime_type: str, name: str) -> bool:
    return str(mime_type or "").strip().lower() == _DOCX_MIME_TYPE or str(name or "").strip().lower().endswith(".docx")


def _is_editable_docx_attachment(item: Any) -> bool:
    if not isinstance(item, dict) or not str(item.get("data_url") or "").strip():
        return False
    return _is_docx_mime_or_name(str(item.get("mime_type") or ""), str(item.get("name") or ""))


def _iter_document_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_paragraph_text(paragraph: Any, replacement: str) -> None:
    """Replace text while retaining paragraph formatting and the first run's character style."""
    first_run_properties = None
    if getattr(paragraph, "runs", None):
        run_properties = paragraph.runs[0]._element.rPr
        if run_properties is not None:
            first_run_properties = deepcopy(run_properties)
    paragraph.text = replacement
    if first_run_properties is not None and paragraph.runs:
        paragraph.runs[0]._element.insert(0, first_run_properties)


def _parse_document_edit_plan(value: str) -> dict[str, Any] | None:
    text = str(value or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.IGNORECASE)
    decoder = json.JSONDecoder()
    start = text.find("{")
    if start < 0:
        return None
    try:
        loaded, _ = decoder.raw_decode(text[start:])
    except (json.JSONDecodeError, TypeError):
        return None
    if not isinstance(loaded, dict) or not isinstance(loaded.get("edits"), list):
        return None
    edits: list[dict[str, str]] = []
    for item in loaded["edits"][:60]:
        if not isinstance(item, dict):
            continue
        target = str(item.get("target") or "").strip()
        replacement = str(item.get("replacement") or "").strip()
        if target and replacement and len(target) <= 20_000 and len(replacement) <= 20_000:
            edits.append({"target": target, "replacement": replacement})
    return {"edits": edits, "summary": str(loaded.get("summary") or "")[:1000]}
